from __future__ import annotations

import argparse
import json
import random
from collections import Counter
from pathlib import Path
from typing import Any

from schemas import QuestionItem
from utils import load_records_from_directory, normalize_text, read_json


def load_eligible(input_dir: Path, minimum_score: int) -> list[QuestionItem]:
    eligible: list[QuestionItem] = []
    for record in load_records_from_directory(input_dir):
        item = QuestionItem.from_dict(record)
        errors = item.validate()
        if errors:
            raise ValueError(f"{item.question_id}: {'; '.join(errors)}")
        if item.review_status == "accepted" and item.quality_score >= minimum_score:
            eligible.append(item)
    return eligible


def load_weights(path: Path | None) -> dict[str, float]:
    if not path:
        return {}
    payload = read_json(path)
    if not isinstance(payload, dict):
        raise ValueError("Chapter weights must be a JSON object")
    weights = {str(key): float(value) for key, value in payload.items()}
    if any(value <= 0 for value in weights.values()):
        raise ValueError("Chapter weights must be positive")
    return weights


def select_questions(
    pool: list[QuestionItem],
    count: int,
    chapter_weights: dict[str, float],
    seed: int,
) -> list[QuestionItem]:
    if len(pool) < count:
        raise ValueError(f"Only {len(pool)} eligible questions are available; requested {count}")
    rng = random.Random(seed)
    candidates = pool[:]
    rng.shuffle(candidates)
    selected: list[QuestionItem] = []
    chapter_counts: Counter[str] = Counter()
    type_counts: Counter[str] = Counter()
    difficulty_counts: Counter[str] = Counter()
    chunk_counts: Counter[str] = Counter()
    concept_keys: set[str] = set()

    all_chapters = sorted({item.chapter_id for item in pool})
    total_weight = sum(chapter_weights.get(chapter, 1.0) for chapter in all_chapters)
    chapter_targets = {
        chapter: count * chapter_weights.get(chapter, 1.0) / total_weight
        for chapter in all_chapters
    }
    basic_target = round(count * 0.8)

    while len(selected) < count:
        best: QuestionItem | None = None
        best_score = float("-inf")
        for item in candidates:
            concept = normalize_text(item.tested_concept)
            target = chapter_targets[item.chapter_id]
            score = (target - chapter_counts[item.chapter_id]) * 4.0
            score += 2.0 / (1 + type_counts[item.question_type])
            score += 2.0 / (1 + chunk_counts[item.source_chunk_id])
            score += item.quality_score / 100.0
            if concept in concept_keys:
                score -= 8.0
            if item.difficulty == "basic":
                if difficulty_counts["basic"] < basic_target:
                    score += 1.5
                else:
                    score -= 1.0
            elif difficulty_counts["basic"] < basic_target:
                score -= 0.5
            if score > best_score:
                best = item
                best_score = score
        if best is None:
            raise RuntimeError("Unable to select enough questions")
        selected.append(best)
        candidates.remove(best)
        chapter_counts[best.chapter_id] += 1
        type_counts[best.question_type] += 1
        difficulty_counts[best.difficulty] += 1
        chunk_counts[best.source_chunk_id] += 1
        concept_keys.add(normalize_text(best.tested_concept))

    rng.shuffle(selected)
    return selected


def write_exam_xlsx(path: Path, questions: list[QuestionItem]) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font
    except ImportError as exc:
        raise RuntimeError("openpyxl is required. Install requirements.txt.") from exc
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "exam"
    headers = ["exam_no", "question_id", "chapter_id", "question_type", "difficulty", "stem", "A", "B", "C", "D", "E"]
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    for number, item in enumerate(questions, start=1):
        sheet.append(
            [
                number,
                item.question_id,
                item.chapter_id,
                item.question_type,
                item.difficulty,
                item.stem,
                item.options["A"],
                item.options["B"],
                item.options["C"],
                item.options["D"],
                item.options["E"],
            ]
        )
    sheet.freeze_panes = "A2"
    sheet.column_dimensions["F"].width = 70
    for column in ("G", "H", "I", "J", "K"):
        sheet.column_dimensions[column].width = 38
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    workbook.save(path)


def write_answer_key(path: Path, questions: list[QuestionItem]) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError as exc:
        raise RuntimeError("openpyxl is required. Install requirements.txt.") from exc
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "answer_key"
    sheet.append(["exam_no", "question_id", "correct_answer", "quality_score", "source_chunk_id"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    for number, item in enumerate(questions, start=1):
        sheet.append([number, item.question_id, item.correct_answer, item.quality_score, item.source_chunk_id])
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def write_docx(path: Path, questions: list[QuestionItem], with_explanations: bool) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Install requirements.txt.") from exc
    document = Document()
    document.add_heading("Nephrology Board Practice Examination", level=0)
    document.add_paragraph(f"Questions: {len(questions)}")
    for number, item in enumerate(questions, start=1):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{number}. {item.stem}")
        run.bold = True
        run.font.size = Pt(11)
        for key in ("A", "B", "C", "D", "E"):
            document.add_paragraph(f"{key}. {item.options[key]}")
        if with_explanations:
            document.add_paragraph(f"Answer: {item.correct_answer}")
            document.add_paragraph(f"Explanation: {item.explanation}")
            document.add_paragraph(f"Source: {item.source_chunk_id} ({item.source_paragraph_range})")
        document.add_paragraph("")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def write_manifest(path: Path, questions: list[QuestionItem], seed: int) -> None:
    payload: dict[str, Any] = {
        "seed": seed,
        "question_count": len(questions),
        "chapter_distribution": dict(Counter(item.chapter_id for item in questions)),
        "question_type_distribution": dict(Counter(item.question_type for item in questions)),
        "difficulty_distribution": dict(Counter(item.difficulty for item in questions)),
        "question_ids": [item.question_id for item in questions],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Select and export a formal exam set.")
    parser.add_argument("--input-dir", type=Path, default=Path("data/final_item_bank"))
    parser.add_argument("--num-questions", type=int, default=100)
    parser.add_argument("--minimum-score", type=int, default=90)
    parser.add_argument("--weights", type=Path, default=Path("data/blueprints/chapter_weights.json"))
    parser.add_argument("--output-dir", type=Path, default=Path("outputs/exam_sets"))
    parser.add_argument("--exam-id", default="exam_set_001")
    parser.add_argument("--seed", type=int, default=20260627)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    pool = load_eligible(args.input_dir, args.minimum_score)
    weights = load_weights(args.weights if args.weights.exists() else None)
    selected = select_questions(pool, args.num_questions, weights, args.seed)
    base = args.output_dir / args.exam_id
    write_exam_xlsx(base.with_suffix(".xlsx"), selected)
    write_answer_key(args.output_dir / f"{args.exam_id}_answer_key.xlsx", selected)
    write_docx(args.output_dir / f"{args.exam_id}_questions_only.docx", selected, False)
    write_docx(args.output_dir / f"{args.exam_id}_with_explanations.docx", selected, True)
    write_manifest(args.output_dir / f"{args.exam_id}_manifest.json", selected, args.seed)
    print(f"selected {len(selected)} questions -> {args.output_dir}")


if __name__ == "__main__":
    main()
